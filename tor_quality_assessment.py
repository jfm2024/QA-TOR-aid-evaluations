from openai import OpenAI
import os
from textstat import flesch_reading_ease
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

# Initialize the OpenAI client
api_key = "YOUR KEY"  
client = OpenAI(api_key=api_key)

# Function to perform the TOR scoring and assessment 
def assess_tor_with_gpt4o(tor_text):
    messages = [
        {
            "role": "system",
            "content": "You are a helpful assistant that performs quality assessments of Terms of Reference documents."
        },
        {
            "role": "user",
            "content": f"""
                    Please assess the following ToR document according to the scoring guide provided, and return feedback in markdown format:

            ### Terms of Reference Quality Assessment Protocol:

            #### 1. Context and Background
            Objective: Evaluate the extent to which the ToR provide relevant background and context of the program being evaluated.
            - **Criteria**:
              - Inclusion of socio-economic, political, or cultural context pertinent to the evaluation.
              - Reference to previous evaluations or related evidence.
            - **Scoring**:
              - 4: Comprehensive context with all relevant information.
              - 3: Adequate but with minor gaps.
              - 2: Limited context with major gaps.
              - 1: Context missing.

            #### 2. Purpose, Rationale, and Evaluation Objectives
            Objective: Assess whether the ToR clearly state the purpose, intended use, and specific objectives of the evaluation.
            - **Criteria**:
              - Clear rationale on why the evaluation is conducted and for whom.
              - Articulation of the ToR's purpose in terms of accountability or learning functions.
              - Specific evaluation objectives that outline what the evaluation aims to accomplish, serving as a foundation for developing relevant questions and methodology.
            - **Scoring**:
              - 4: Fully clear and detailed purpose, rationale, and objectives.
              - 3: Mostly clear with minor gaps.
              - 2: Described but with significant gaps.
              - 1: Not stated or unclear.

            #### 3. Evaluation Scope
            Objective: Ensure the ToR clearly outline the scope, including geographic, temporal, and thematic boundaries of the evaluation.
            - **Criteria**:
              - Clear definition of geographic, temporal, and thematic scope, delineating the focus areas of the evaluation.
            - **Scoring**:
              - 4: Fully defined scope with relevant details.
              - 3: Mostly clear scope, minor gaps in detail.
              - 2: Limited description of scope with significant gaps.
              - 1: Scope not provided or unclear.
            - **Additional Commentary (non-scored)**: Please add any relevant commentary on whether the ToR reference specific evaluation criteria (e.g., OECD-DAC criteria), as this can enhance the evaluation’s alignment with standard frameworks.

            #### 4. Evaluation Questions and Relevance
            Objective: Determine if the ToR include tailored and relevant evaluation questions.
            - **Criteria**:
              - Questions are specific, aligned with the evaluation objectives (outlined under Purpose), and relevant to the intended use.
            - **Scoring**:
              - 4: Evaluation questions are specific, relevant, and customized.
              - 3: Mostly specific and relevant, minor gaps.
              - 2: Evaluation questions lack relevance or specificity.
              - 1: No clear questions.

            #### 5. Methodology and Feasibility
            Objective: Assess if the ToR provide a feasible methodological approach considering the evaluation questions, given the timeframe and resources.
            - **Criteria**:
              - Proposed methodology is clear, suitable for addressing the evaluation questions, and feasible within the set timeline.
            - **Scoring**:
              - 4: Clear, feasible methodology aligned with evaluation questions.
              - 3: Feasible with minor adjustments needed.
              - 2: Major feasibility issues.
              - 1: Unfeasible or not described.

            #### 6. Roles and Responsibilities
            Objective: Evaluate whether the ToR clearly outline roles, responsibilities, and references to quality assurance steps.
            - **Criteria**:
              - Clear definition of roles for consultants, team members, and commissioners.
              - References to quality assurance steps in the evaluation.
            - **Scoring**:
              - 4: All roles and quality assurance steps clearly defined.
              - 3: Mostly clear with minor ambiguities.
              - 2: Described but with major gaps.
              - 1: Not defined or unclear.
            - **Additional Commentary (non-scored)**: Please include commentary on process elements in the ToR, such as evaluation phases (e.g., inception, data collection, reporting) and other process-related instructions, as these insights can highlight procedural guidance in the ToR.

            #### 7. Ethics, Risk, and Mitigation Strategies
            Objective: Evaluate whether the ToR address ethical considerations, potential risks, and mitigation strategies relevant to the evaluation.
            - **Criteria**:
              - Mention of ethical standards or safeguards, with relevant considerations such as human rights, gender, or environmental impact.
              - Identification of key risks (e.g., data availability, stakeholder engagement, logistical challenges).
              - Mitigation strategies that are practical and aligned with the evaluation's scope and timeline.
            - **Scoring**:
              - 4: Comprehensive ethical considerations and risk assessment with robust mitigation strategies.
              - 3: Adequate ethics and risk assessment with minor gaps in mitigation plans.
              - 2: Limited ethical considerations or risk assessment with insufficient mitigation strategies.
              - 1: Ethics or risk assessment not included or poorly addressed.

            #### 8. Deliverables
            Objective: Ensure the ToR clearly outline all required deliverables, including reports and datasets, to support transparency and replicability.
            - **Criteria**:
              - Specifies all essential deliverables, including an inception report, a draft report, and a final report, with clear guidelines on content, format, quality standards, and submission timelines.
              - Includes a requirement for datasets generated or used during the evaluation, detailing standards for documentation and organization to facilitate replicability and future use.
              - Lists any additional deliverables (e.g., executive summaries, presentations) and provides expectations for these.
            - **Scoring**:
              - 4: Comprehensive list of deliverables, including datasets with complete guidance on format, standards, and timelines for each deliverable.
              - 3: Deliverables and expectations are mostly clear, including datasets, but minor gaps in guidance or format expectations exist.
              - 2: Limited information on deliverables, including datasets, with significant gaps in format or expectations.
              - 1: Deliverables, including datasets, are not clearly defined or missing from the ToR.

            ### Document Text:
            '{tor_text}'

            Provide your evaluation in markdown, including scores, comments, and any additional commentary specified in each section.
            """
        }
    ]

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=messages,
        temperature=0.7,
        max_tokens=1000
    )
    return response.choices[0].message.content


# Function to get evaluation team feedback
def review_as_evaluation_team(tor_text):
    messages = [
        {
            "role": "system",
            "content": "You are an evaluator focusing on practical and operational clarity for the evaluation team commissioned to execute this assignment."
        },
        {
            "role": "user",
            "content": f"""
            From the perspective of the evaluation team tasked with implementing this ToR, assess the document's clarity, feasibility, and guidance. Consider the following points in your review:
            - Does the ToR clearly define the team’s tasks and objectives?
            - Is the proposed methodology clear?
            - Are roles, responsibilities, and deliverable expectations clearly outlined for smooth execution?
            - Highlight any ambiguities or gaps that could impact the team’s work.
            
            Provide your feedback in 3-5 sentences.Write in the style of an evaluation team member.

            ### Terms of Reference (ToR):
            {tor_text}
            """
        }
    ]

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=messages,
        temperature=0.7,
        max_tokens=400
    )
    return response.choices[0].message.content

# Function to get Norwegian development aid bureaucrat feedback
def review_as_norwegian_bureaucrat(tor_text):
    messages = [
        {
            "role": "system",
            "content": "You are an evaluator analyzing this ToR for clarity and policy alignment, as it would be reviewed by a bureaucrat in the Norwegian development aid system."
        },
        {
            "role": "user",
            "content": f"""
From the perspective of a bureaucrat in the Norwegian development aid system, assess the ToR’s clarity and alignment with policy standards. Consider the following points in your review:
- Does the ToR clearly outline objectives?
- Is the rationale behind the evaluation clear?
- Are the the key concepts understandable?
- Highlight any areas that could be confusing.

Provide your feedback in 3-5 sentences. Write in the style of a bureaucrat in the Norwegian development aid system.

### Terms of Reference (ToR):
{tor_text}
"""
        }
    ]

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=messages,
        temperature=0.7,
        max_tokens=400
    )
    return response.choices[0].message.content

# Function to check readability
def check_readability(text):
    return flesch_reading_ease(text)

# Function to generate a concise summary of the combined feedback
def generate_summary(gpt4o_feedback, team_feedback, bureaucrat_feedback, readability_score):
    summary_prompt = f"""
    Based on the following feedback from three evaluators, and the Readability Score, provide a very concise 3-5 sentence summary highlighting the main strengths and areas for improvement of the Terms of Reference (ToR) document.

    ### GPT-4o Scoring and Feedback:
    {gpt4o_feedback}

    ### Evaluation Team Perspective:
    {team_feedback}

    ### Readability Score:
    {readability_score}
    
    ### Norwegian Bureaucrat Perspective:
    {bureaucrat_feedback}
    """
    
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": summary_prompt}],
        temperature=0.7,
        max_tokens=250  
    )

    return response.choices[0].message.content


def generate_final_report(tor_text):
    # Get feedback from each function
    gpt4o_feedback = assess_tor_with_gpt4o(tor_text).replace("```markdown", "").replace("---", "").replace("### ", "").replace("## ", "").replace("**", "").replace("# ", "").replace("```", "").strip()
    team_feedback = review_as_evaluation_team(tor_text).replace("```markdown", "").replace("---", "").replace("### ", "").replace("## ", "").replace("**", "").replace("# ", "").replace("```", "").strip()
    bureaucrat_feedback = review_as_norwegian_bureaucrat(tor_text).replace("```markdown", "").replace("---", "").replace("### ", "").replace("## ", "").replace("**", "").replace("# ", "").replace("```", "").strip()
    
    # Calculate readability score
    readability_score = check_readability(tor_text)
    
    # Generate summary based on combined feedback
    summary = generate_summary(gpt4o_feedback, team_feedback, bureaucrat_feedback, readability_score).replace("```markdown", "").replace("---", "").replace("### ", "").replace("## ", "").replace("**", "").strip()
    
    # Create a new Word document
    doc = Document()
    
    title = doc.add_heading("Terms of Reference Quality Assurance Report", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def set_line_spacing(paragraph):
        paragraph.paragraph_format.line_spacing = 1.5
    
    # Section 1: Quality Assessment including Scoring
    doc.add_heading("1. Quality Assessment including Scoring", level=2)
    scoring_paragraph = doc.add_paragraph(gpt4o_feedback)
    set_line_spacing(scoring_paragraph)
    doc.add_paragraph().add_run().add_break()  
    
    # Section 2: Evaluation Team Perspective
    doc.add_heading("2. Evaluation Team Perspective", level=2)
    team_paragraph = doc.add_paragraph(team_feedback)
    set_line_spacing(team_paragraph)
    doc.add_paragraph().add_run().add_break()  
    
    # Section 3: Norwegian Development Aid Bureaucrat Perspective
    doc.add_heading("3. Norwegian Development Aid Bureaucrat Perspective", level=2)
    bureaucrat_paragraph = doc.add_paragraph(bureaucrat_feedback)
    set_line_spacing(bureaucrat_paragraph)
    doc.add_paragraph().add_run().add_break()  
    
    # Section 4: Readability Score
    doc.add_heading("4. Readability Score", level=2)
    readability_paragraph = doc.add_paragraph(f"Readability Score: {readability_score}")
    set_line_spacing(readability_paragraph)
    
    # Explanation of Readability Score
    doc.add_heading("Readability Score Explanation", level=3)
    explanation_text = (
        "The Readability Score, based on the Flesch Reading Ease formula, indicates how easy the Terms of Reference (ToR) document is to read. "
        "Scores range from 0 to 100, with higher scores meaning easier readability:\n\n"
        "- 90-100: Very easy to read, suitable for young readers.\n"
        "- 60-70: Standard readability, generally understandable by teenagers.\n"
        "- 30-50: Difficult to read, best suited for adults with college-level reading skills.\n"
        "- Below 30: Very complex, suitable for advanced readers.\n\n"
        "This score can help identify areas where simplifying language or structure might improve accessibility."
    )
    explanation_paragraph = doc.add_paragraph(explanation_text)
    set_line_spacing(explanation_paragraph)
    
    # Section 5: Summary
    doc.add_heading("Summary", level=2)
    summary_paragraph = doc.add_paragraph(summary)
    set_line_spacing(summary_paragraph)
    
    # Save the document
    doc.save("ToR_Quality_Assurance_Report.docx")
    print("Report saved as ToR_Quality_Assurance_Report.docx")
    
# Collect TOR input from the user
print("Please enter your TOR text. Type 'END' on a new line to finish input.")
user_input = []
while True:
    line = input()
    if line.strip().upper() == "END":
        break
    user_input.append(line)

# Join all lines into a single string
document_text = "\n".join(user_input)

# Run the final report generation
generate_final_report(document_text)